#!/usr/bin/env python
# coding: utf-8

# # 1. Automated Data Profiling In Python

# **Author : Anandakumar Varatharajah**
# <br>
# ***http://www.analyticsinsights.ninja***

# Version   : 0.17   
# Date      : 14 July 2019  
# License   : MIT License

# The main objective of this notebook is **only** to understand raw data profile. i.e. data type, min & max values, ranges, unique values, etc.  
# In consequent notebooks we will explore further on how to make decisions to make the data tidy and perform the data transformations based on the understanding of the data profile.
# <br>
# The code is largely kept generic so that it could be used with any shape of data. 

# # The Game Changer - Data Profile Dataframe (DPD)

# The game changer for exploratory data analysis is the final ***Data Profile Dataframe*** that is generated which combines ***all*** the information required to inform data cleaning, tidy data and optimisations (memory and processing) decisions.  
# Instead of using various Pandas commands at different instances and going back and forth to cross refer information, Data Profile Dataframe brings all information into a single dataframe. This will be very useful when reviewing the data profile with the business subject matter or other team members as all information related to data profile is in a single easy to understand format.
# 
# ![image.png](https://raw.githubusercontent.com/AnalyticsInsightsNinja/Python_TidyData/master/SAMPLE_FULL_DPD_Image_MSWORD.PNG)
# 

# Understanding the data is **the critical step** in preparing the data to be used for analytics. As many experts will point out the data preparation and transforming the data into a tidy format takes about 80% of the effort in any data analytics or data analysis project.<br>
# ***Understanding the data requires good understanding of the domain and/or access to a subject matter expert (SME) to help make decisions about data quality and data usage:***
# * What are the columns and what do they mean?
# * How to interpret each columns and possible values of a column?
# * Should the columns be renamed (and cleaned e.g. trim)?
# * Are there columns that may have similar information that could be dropped in favour of one master column?
# * Can columns with no values (or all empty) be dropped?
# * Can columns which have more than certain threshold of blank values be dropped?
# * How can the missing values be filled and can it be filled meaningfully?
# * Can rows that have missing values for certain columns or combination of columns be dropped? i.e. the row is meaningless wihtout those values.
# * Can the numeric data type columns be converted / down casted to optimise memory usage based on the data values?
#     - or will there be outliers possibly in future data sets that we cannot do this?
#     - can the min and max values be used to determine the lowest possible data type?
# * Can some string/object columns be converted to Category types?
#     - based on count of unique values
# * Can any columns be discarded that may not be required for analytics?

# # Environment setup

# It is recommended best practice to document the execution environment.  
# e.g. When the initial version of this notebook was developed in Azure Notebooks (Jupyter) the environment was documented in the code. When the notebook was exported to local PC JupyterLab and then imported back into Azure Notebook, the Kernal changed to an older version and some code did not work. Having the initital versions documented in comments saved a lot of effort in trying to understand what went wrong.
# 

# In[ ]:


# Get the date of execution
import datetime
date_generated = datetime.datetime.now()


# In[ ]:


from platform import python_version 
# use python_version() to get the version. This is used in the final DPD HTML
# 3.6.6 in Azure Notebooks in April 2019


# In[ ]:


import pandas as pd
# use pd.__version__ to get the pandas version. This is used in the final DPD HTML
# Pandas version   0.22.0 in Azure Notebooks in April 2019

# set maximum number of columns to display in notebook
pd.set_option('display.max_columns', 250)

# To check whether a column is numeric type
from pandas.api.types import is_numeric_dtype

# To check whether a column is object/string type
from pandas.api.types import is_string_dtype


# In[ ]:


import numpy as np


# In[ ]:


# Import the graph packages
import matplotlib.pyplot as plt
plt.rcParams.update({'figure.max_open_warning': 0})
get_ipython().run_line_magic('matplotlib', 'inline')

import seaborn as sns
# Seabotn version   0.9.0 in Azure Notebooks in April 2019
# use sns.__version__ to get the pandas version. This is used in the final DPD HTML


# In[ ]:


# This library is required to generate the MS Word document
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH  #used to align str(number) in cells 


# # Raw data file exploration

# The raw data file used in this notebook has been derived from the Sales Products csv file from IBM Analytics Community and has been modified to include untidy data for the purposes of this data exploration work.  
# The raw data should be in a format that can be laoded into pandas. i.e. if there are any rows need to be skipped,  column headers mapped, etc. should be handle in the pandas.read code block.

# In[ ]:


# Download data file from Github site using curl and save it to local disk   -o "filename"
get_ipython().system('curl -o "mydataset.csv" "https://raw.githubusercontent.com/AnalyticsInsightsNinja/Sample_Analytics_Data/master/titanic.csv" ')
# Data file to be loaded
raw_data_file = "mydataset.csv"


# In[ ]:


# Use Pandas to load the data file into a dataframe
try:
    df = pd.read_csv(raw_data_file, thousands=',', float_precision=2)
except:
    print("Error: Data file not found!")


# **Note:** If the raw data is a big data file of several GB's in size it may not be possible to load the the whole file into memory. One possibility is using 'pandas pyspark'.<br>
# Other options to load data incrementally and optimise the data by converting data types will be demonstrated in a seperate notebook.

# In[ ]:


# Sample raw data rows from dataset
df.sample(5).round(2)


# # Memory Usage Analysis

# In[ ]:


# Check whether the file is obatined from url.
# If from url, then skip file size in disk check

if "http" in raw_data_file:
    file_size = float('nan')
else:
    # Calculating file size (in MB) on disk
    import os

    file_size = (os.stat(raw_data_file).st_size / 1024 **2)
    #This is used in the DPD HTML


# In[ ]:


# Calculate dataset size in memory (MB)
df_mem = df.memory_usage(deep=True).sum() / 1024**2
#This is used in the DPD HTML


# In[ ]:


# Calclulate dataset size increase in memory (MB)
sz_increase = ((df_mem - file_size) / file_size)
#This is used in the DPD HTML


# In[ ]:


# Plot the memory usage 
# Create a dictionary from the variables and convert to Pandas DataFrame
# Use DataFrame's ploting capabilities
raw_data_dict = {"File on disk":file_size, "Dataset in memroy": df_mem}
raw_data_plot = pd.DataFrame.from_dict(raw_data_dict, orient='index').reset_index()

# Pandas DataFrame plot
raw_data_plot.plot(kind='bar',                   x="index" ,                   y=0,                    legend=False,                    title='Data size increase from disk to memory')
# plt.subplots_adjust(wspace=0.4, hspace=0.35)
plt.xticks(rotation=0)

# Save the figure
plt.savefig('fig_df_tot_memory.png', dpi=50)
plt.close('all')


# In[ ]:


# Get memory used by each column in the raw data dataset in MB
# This will be later merged with the DPD
mem_used_dtypes = pd.DataFrame(df.memory_usage(deep=True) / 1024**2)

# Rename column
mem_used_dtypes.rename(columns={ 0:'memory'}, inplace=True)

# Drop index memory usage since this is not required when merging with Data Quality Dataframe
mem_used_dtypes.drop('Index', axis=0, inplace=True) 


# # Constructing The Data Profile Dataframe (DPD) - The Game Changer 

# In[ ]:


# Number of rows of the DPD will be the count of columns in the raw date dataframe
# Since it there will be one row for each column
no_of_rows = len(df.columns)


# Constructing the data_qlt_df dataframe and pre-assigning and columns
# Pre-assigning the number of rows the dataframe would have is memory and processing efficient
# This is a better approach than continuous append or concat operation to dataframe

data_qlt_df = pd.DataFrame(index=np.arange(0, no_of_rows),                             columns=('column_name', 'col_data_type', 'col_memory','non_null_values',                                      'unique_values_count', 'column_dtype')
                          )


# Add rows to the data_qlt_df dataframe
for ind, cols in enumerate(df.columns):
    # Count of unique values in the column
    col_unique_count = df[cols].nunique()
    
    data_qlt_df.loc[ind] = [cols,                             df[cols].dtype,                             mem_used_dtypes['memory'][ind],                             df[cols].count(),                             col_unique_count,                             cols + '~'+ str(df[cols].dtype)
                            ]


# In[ ]:


# Use describe() to get column stats of raw dataframe
# This will be merged with the DPD
raw_num_df = df.describe().T.round(2)


# In[ ]:


#----- Key Step ---------------
# Merging the df.describe() output with rest of the info to create a single Data Profile Dataframe
data_qlt_df = pd.merge(data_qlt_df, raw_num_df, how='left', left_on='column_name', right_index=True)


# In[ ]:


# Calculate percentage of non-null values over total number of values
data_qlt_df['%_of_non_nulls'] = (data_qlt_df['non_null_values']/df.shape[0])*100

# Calculate null values for the column
data_qlt_df['null_values'] = df.shape[0] - data_qlt_df['non_null_values']

# Calculate percentage of null values over total number of values
data_qlt_df['%_of_nulls'] = 100 - data_qlt_df['%_of_non_nulls']

# Calculate percentage of each column memory usage compared to total memory used by raw data datframe
data_qlt_df['%_of_total_memory'] = data_qlt_df['col_memory'] / data_qlt_df['col_memory'].sum() * 100

# Calculate the total memory used by a given group of data type
# See Notes section at the bottom of this notebook for advatages of using 'transform' function with group_by
data_qlt_df["dtype_total"] = data_qlt_df.groupby('col_data_type')["col_memory"].transform('sum')

# Calculate the percentage memory used by each column data type compared to the total memory used by the group of data type
# the above can be merged to one calculation if we do not need the total as separate column
#data_qlt_df["%_of_dtype_mem2"] = data_qlt_df["Dtype Memory"] / (data_qlt_df.groupby('Data Type')["Dtype Memory"].transform('sum')) * 100
data_qlt_df["%_of_dtype_mem"] = data_qlt_df["col_memory"] / data_qlt_df["dtype_total"] * 100

# Calculate the percentage memory used by each group of data type of the total memory used by dataset
data_qlt_df["dtype_%_total_mem"] = data_qlt_df["dtype_total"] / df_mem * 100

# Calculate the count of each data type
data_qlt_df["dtype_count"] = data_qlt_df.groupby('col_data_type')["col_data_type"].transform('count')

# Calculate the total count of column values
data_qlt_df["count"] = data_qlt_df['null_values'] + data_qlt_df['non_null_values']


# In[ ]:


# Reorder the Data Profile Dataframe columns
data_qlt_df = data_qlt_df[
                            ['column_name', 'col_data_type', 'col_memory', '%_of_dtype_mem', '%_of_total_memory',\
                             'dtype_count', 'dtype_total', 'dtype_%_total_mem', 'non_null_values', '%_of_non_nulls',\
                             'null_values', '%_of_nulls', 'unique_values_count', 'count', 'mean', 'std', 'min', '25%',\
                             '50%', '75%', 'max']
                         ]


# **The above data quality data frame summarises all information required for making data quality decisions.**  
# Though there are info() and describe() methods to do these, having all the relvant information in one dataframe makes the data quality exploration much easier. This dataframe can be used for summarising information and for plotting to ehnace the ease of Data Understanding effort.

# # Plot Memory Usage Analysis

# In[ ]:


# Plot count of column data types and memory used by each datatype
plt_dtype = data_qlt_df.groupby('col_data_type')['dtype_count', 'dtype_total', 'dtype_%_total_mem'].last().sort_values(by='dtype_count')

fig1, (ax, ax2) = plt.subplots(ncols=2, figsize=(10,5))
plt.subplots_adjust(wspace=0.4, hspace=0.35, bottom=0.20)

plt_dtype.plot(kind='bar', y='dtype_count',  use_index=True, legend=False, ax=ax, title='Count of columns by data type')

plt_dtype.plot(kind='bar', y='dtype_total',  use_index=True, legend=False, ax=ax2, title='Memory used by data type')

fig1.savefig("fig_cols_memory.png", dpi=50)
plt.close('all')


# In[ ]:


# Memory used by columns of raw data dataframe
fig2, ax = plt.subplots(ncols=1, figsize=(15,5))
plt.subplots_adjust(wspace=0.4, hspace=0.35, bottom=0.30)

# Memory used by object data type
(data_qlt_df[data_qlt_df['col_data_type'] == 'object']
 .sort_values(by='col_memory', ascending=False)
 .plot(kind="bar", 
       x="column_name", 
       y="col_memory", 
       title="Memory (MB) usage by columns of object data type",
      legend=False, ax=ax)
)
plt.xticks(rotation=35)
fig2.savefig("fig_object_cols_memory.png", dpi=50)
plt.close('all')

# Memory used by non-object data type
fig2, ax1 = plt.subplots(ncols=1, figsize=(15,5))
plt.subplots_adjust(wspace=0.4, hspace=0.35, bottom=0.30)

(data_qlt_df[data_qlt_df['col_data_type'] != 'object']
 .sort_values(by='col_memory', ascending=False)
 .plot(kind="bar", 
       x="column_name", 
       y="col_memory", 
       title="Memory (MB) usage by columns of non-object data type",
      legend=False, ax=ax1)
)
plt.xticks(rotation=35)

fig2.savefig("fig_non_object_cols_memory.png", dpi=50)
plt.close('all')


# # Generate data profile graphs for 'numerical' columns

# In[ ]:


import numpy as np
from matplotlib.patches import Rectangle

# Get the list of numeric columns from raw dataframe
# need this: from pandas.api.types import is_numeric_dtype
# get numeric columns which are not empty
num_cols = [cols for cols in df.columns if is_numeric_dtype(df[cols]) and len(df[cols].dropna())>0]

iter_len = len(num_cols)

# For each numeric column in the list
for x, col_name in enumerate(num_cols):
    print(x+1, " of ", iter_len, " completed   ",  col_name)
    
    # Create a copy of the column values without nulls or NA
    no_null_col = df[col_name].dropna()
    
    
    # Calculate the 95 percentile of the values
    q25 = np.percentile(no_null_col, 25)
    q75 = np.percentile(no_null_col, 75)    
    q95 = np.percentile(no_null_col, 95)
    
    # Plot the graphs
    fig3 = plt.figure(figsize=(20,15))
    fig3.suptitle("Profile of column  " + col_name, fontsize=25)  #Title for the whole figure
    plt.subplots_adjust(wspace=0.4, hspace=0.35)

    ax1 = fig3.add_subplot(2,3,1)
    ax1.set_title("Box plot for all the values", fontsize=20)
    plt.setp(ax1.get_xticklabels(), ha="right", rotation=35)
    plt.setp(ax1.get_yticklabels(), ha="right", fontsize=15)
    ax1.boxplot(no_null_col)

    ax1 = fig3.add_subplot(2,3,2)
    ax1.set_title("Distribution of all values", fontsize=20)
    plt.setp(ax1.get_xticklabels(), ha="right", rotation=35, fontsize=15)
    plt.setp(ax1.get_yticklabels(), ha="right", fontsize=15)
    ax1.hist(no_null_col)

    ax1 = fig3.add_subplot(2,3,3)
    ax1.set_title("Boxplot for quartiles (all values)", fontsize=20)
    if len(no_null_col.value_counts()) >= 4:
        df[u'quartiles'] = pd.qcut(
                        df[col_name],
                        4, duplicates='drop')
        df.boxplot(column= col_name, by=u'quartiles', ax = ax1)
    plt.setp(ax1.get_xticklabels(), ha="right", rotation=35, fontsize=15)
    plt.setp(ax1.get_yticklabels(), ha="right", fontsize=15)

    ax1 = fig3.add_subplot(2,3,4)
    ax1.set_title("Box plot without outliers", fontsize=20)
    plt.setp(ax1.get_xticklabels(), ha="right", rotation=35, fontsize=15)
    plt.setp(ax1.get_yticklabels(), ha="right", fontsize=15)
    ax1.boxplot(no_null_col, showfliers=False)

    ax1 = fig3.add_subplot(2,3,5)
    ax1.set_title("Violin plot (<95% percentile)", fontsize=20)
    plt.setp(ax1.get_xticklabels(), ha="right", rotation=35, fontsize=15)
    plt.setp(ax1.get_yticklabels(), ha="right", fontsize=15)
    ax1.violinplot(no_null_col[no_null_col <= q95])

    
    #Histogram with bin ranges, counts and percentile color
    ax1 = fig3.add_subplot(2,3,6)
    ax1.set_title("Histogram (<95% percentile)", fontsize=20)
    plt.setp(ax1.get_xticklabels(), ha="right", rotation=35, fontsize=15)
    plt.setp(ax1.get_yticklabels(), ha="right", fontsize=15)

    # Take only the data less than 95 percentile
    data = no_null_col[no_null_col <= q95]

    # Colours for different percentiles
    perc_25_colour = 'gold'
    perc_50_colour = 'mediumaquamarine'
    perc_75_colour = 'deepskyblue'
    perc_95_colour = 'peachpuff'

    '''
    counts  = numpy.ndarray of count of data ponts for each bin/column in the histogram
    bins    = numpy.ndarray of bin edge/range values
    patches = a list of Patch objects.
            each Patch object contains a Rectnagle object. 
            e.g. Rectangle(xy=(-2.51953, 0), width=0.501013, height=3, angle=0)
    '''
    counts, bins, patches = ax1.hist(data, bins=10, facecolor=perc_50_colour, edgecolor='gray')

    # Set the ticks to be at the edges of the bins.
    ax1.set_xticks(bins.round(2))
    plt.xticks(rotation=70, fontsize=15)

    # Change the colors of bars at the edges
    for patch, leftside, rightside in zip(patches, bins[:-1], bins[1:]):
        if rightside < q25:
            patch.set_facecolor(perc_25_colour)
        elif leftside > q95:
            patch.set_facecolor(perc_95_colour)
        elif leftside > q75:
            patch.set_facecolor(perc_75_colour)

    # Calculate bar centre to display the count of data points and %
    bin_x_centers = 0.5 * np.diff(bins) + bins[:-1]
    bin_y_centers = ax1.get_yticks()[1] * 0.25

    # Display the the count of data points and % for each bar in histogram
    for i in range(len(bins)-1):
        bin_label = "{0:,}".format(counts[i]) + "  ({0:,.2f}%)".format((counts[i]/counts.sum())*100)
        plt.text(bin_x_centers[i], bin_y_centers, bin_label, rotation=90, rotation_mode='anchor')

    #create legend
    handles = [Rectangle((0,0),1,1,color=c,ec="k") for c in [perc_25_colour, perc_50_colour, perc_75_colour, perc_95_colour]]
    labels= ["0-25 Percentile","25-50 Percentile", "50-75 Percentile", ">95 Percentile"]
    plt.legend(handles, labels, bbox_to_anchor=(0.5, 0., 0.85, 0.99))
    

    fig3.suptitle("Profile of column  " + col_name, fontsize=25)  #Title for the whole figure
    fig_name = 'fig_' + col_name
    fig3.savefig(fig_name, dpi=50)
    plt.close('all')
    
#     plt.show()

df.drop(u'quartiles', axis=1, inplace=True)


# # Generate data profile graphs for 'object' columns

# In[ ]:


# Get the list of object columns from raw dataframe
# get object columns which are not empty
obj_cols = [cols for cols in df.columns if is_string_dtype(df[cols]) and len(df[cols].dropna())>0]

iter_len = len(obj_cols)


# For each object column in the list
for x, col_name in enumerate(obj_cols):
    print(x+1, " of ", iter_len, " completed   ",  col_name)
    
    # Create a copy of the column values without nulls or NA
    no_null_col = df[col_name].dropna()

    values_freq_threshold = 25
    col_unique_count = df[col_name].nunique()
    
    # If unique values count is below the threshold value then store the details of unique values
    col_unique_vals = df[col_name].value_counts(normalize=True, sort=True)
    
    # Plot the graphs
    fig4 = plt.figure(figsize=(20,7))
    fig4.suptitle("Profile of column  " + col_name, fontsize=25)  #Title for the whole figure
    plt.subplots_adjust(wspace=0.4, hspace=0.35, bottom=0.35)

    ax1 = fig4.add_subplot(1,1,1)
    ax1.set_title("Bar chart for top 25 values", fontsize=20)
    plt.setp(ax1.get_xticklabels(), ha="right", rotation=45, fontsize=15)
    plt.setp(ax1.get_yticklabels(), ha="right", fontsize=15)
    
    col_unique_vals.head(values_freq_threshold).sort_values(ascending=False).plot.bar()
    plt.xticks(rotation=75)
    for p in ax1.patches:
        ax1.annotate(str(round(p.get_height(),2)), (p.get_x() * 1.005, p.get_height() * 1.005), fontsize=15)
    
    fig4.suptitle("Profile of column  " + col_name, fontsize=25)  #Title for the whole figure
    fig_name = 'fig_' + col_name
    fig4.savefig(fig_name, dpi= 50)

    plt.close('all')
#     plt.show()


# # Candidate columns for Category type

# Analysing how many unique values an 'object' column has will be useful to detrmine which columns are good candidates for *Categorical* data type. In combination with the total memory used by 'object' data type and each 'object' data type column, decisions can be made on converting them Category type.

# In[ ]:


# Create a df and a column for % of memory by each object column
cardn_df = data_qlt_df[data_qlt_df['col_data_type'] == 'object'][['column_name', 'col_memory', '%_of_dtype_mem', '%_of_total_memory', 'unique_values_count']]

cardn_df = cardn_df.sort_values('unique_values_count')


# # Candidate columns for down casting type

# In[ ]:


# Create a df and a column for % of memory by each object column
num_cardn_df = data_qlt_df[data_qlt_df['col_data_type'] != 'object'][['column_name', 'col_memory', '%_of_dtype_mem', '%_of_total_memory', 'unique_values_count']]

num_cardn_df = num_cardn_df.sort_values('unique_values_count')


# # Columns with high percentage of null values

# In[ ]:


# The empty values threshold can be set to a lower/higher value depending on the size of the data sets 
threshold_perc = 0.75
col_vals_threshold = df.shape[0] * threshold_perc


# In[ ]:


null_vals_df = data_qlt_df[data_qlt_df['non_null_values'] < col_vals_threshold][['column_name', 'col_data_type', 'col_memory', 'non_null_values', '%_of_non_nulls', 'null_values', '%_of_nulls']]

# .style.format({'dtype_memory': "{:,.2f}", 'non_null_values': "{:,.2f}", '%_of_non_nulls': "{:,.2f}", 'null_values': "{:,.2f}", '%_of_nulls': "{:,.2f}",  'unique_values_count': "{:,.2f}"})


# # Generate the Correlation plot

# In[ ]:


f, ax = plt.subplots(figsize=(15, 10))
plt.subplots_adjust(bottom=0.35)
plt.autoscale()

corr_data = df.corr()
sns.heatmap(corr_data,
            mask=np.zeros_like(corr_data, dtype=np.bool), 
            cmap=sns.diverging_palette(20, 220, as_cmap=True),
            vmin=-1, vmax=1,
            square=True, 
            ax=ax)

fig_name = 'fig_cor_plot.png'
f.savefig(fig_name,  dpi=70)
# plt.show()
plt.close('all')


# ### Importing the image for the document

# In[ ]:


import requests

image_url = "https://raw.githubusercontent.com/AnalyticsInsightsNinja/Python_TidyData/master/SAMPLE_FULL_DPD_Image_MSWORD.PNG"

Picture_request = requests.get(image_url)
if Picture_request.status_code == 200:
    with open("msword_output.jpg", 'wb') as f:
        f.write(Picture_request.content)


# # Construct the MS Word document

# In[ ]:


# Make sure you have the docx package and it is imported
# see the environment setup section

#Create Document object
document = Document()

# Add Title
document.add_heading('Data Profile Dataframe - Notebook v0.17 - 14 July 2019', 0)
document.add_heading(raw_data_file, 0)

# COver page paragraph
p = document.add_paragraph('The main objective of this notebook is ')
p.add_run('only').bold = True
p.add_run(' to understand raw data profile. i.e. data type, min & max values, ranges, unique values, etc.')
p = document.add_paragraph('In consequent notebooks we will explore further on how to make decisions to make the data tidy and perform the data transformations based on the understanding of the data profile.')
p = document.add_paragraph('')
p.add_run('The code is largely kept generic so that it could be used with any shape of data.').italic = True


# In[ ]:


# Page 2
document.add_page_break()
# Heading 1
document.add_heading('The Game Changer - Data Profile Dataframe (DPD)', level=1)
p = document.add_paragraph('The game changer for exploratory data analysis is the final')
p.add_run(' Data Profile Dataframe').bold = True
p.add_run(' that is generated which combines ')
p.add_run('all').bold = True
p.add_run(' the information required to inform data cleaning, tidy data and optimisations (memory and processing) decisions. Instead of using various Pandas commands at different instances and going back and forth to cross refer information, Data Profile Dataframe brings all information into a single dataframe. This will be very useful when reviewing the data profile with the business subject matter or other team members as all information related to data profile is in a single easy to understand format.')

document.add_picture('msword_output.jpg', height=Inches(4), width=Inches(4))

document.add_page_break()
p = document.add_paragraph('Understanding the data is ')
p.add_run('the critical step').bold = True
p.add_run(' in preparing the data to be used for analytics. As many experts will point out the data preparation and transforming the data into a tidy format takes about 80% of the effort in any data analytics or data analysis project.')
p = document.add_paragraph('')
p.add_run('Understanding the data requires good understanding of the domain and/or access to a subjectmatter expert (SME) to help make decisions about data quality and data usage:').bold = True

document.add_paragraph(
    'What are the columns and what do they mean?', style='List Bullet'
)
document.add_paragraph(
    'How to interpret each columns and possible values of a column?', style='List Bullet'
)
document.add_paragraph(
    'Should the columns be renamed (and cleaned e.g. trim)?', style='List Bullet'
)
document.add_paragraph(
    'Are there columns that may have similar information that could be dropped in favour of one master column?', style='List Bullet'
)
document.add_paragraph(
    'Can columns with no values (or all empty) be dropped?', style='List Bullet'
)
document.add_paragraph(
    'Can columns which have more than certain threshold of blank values be dropped?', style='List Bullet'
)
document.add_paragraph(
    'Can rows that have missing values for certain columns or combination of columns be dropped?', style='List Bullet'
)
document.add_paragraph(
    'i.e. the row is meaningless wihtout those values.', style='List Continue'
)
document.add_paragraph(
    'Can the numeric data type columns be converted / down casted to optimise memory usage based on the data values?', style='List Bullet'
)
document.add_paragraph(
    'or will there be outliers possibly in future data sets that we cannot do this?', style='List Bullet 2'
)
document.add_paragraph(
    'Can the min and max values be used to determine the lowest possible data type?', style='List Bullet 2'
)
document.add_paragraph(
    'Can some string/object columns be converted to Category types?', style='List Bullet'
)
document.add_paragraph(
    'based on count of unique values', style='List Bullet 2'
)
document.add_paragraph(
    'Can any columns be discarded that may not be required for analytics?', style='List Bullet'
)


# # Word - Data profile summary

# In[ ]:


document.add_page_break()
document.add_heading('Columns Data Profile Summary', 0)


# In[ ]:


# Page 4
p = document.add_paragraph(' ')

# Heading 1
document.add_heading('Dataset shape', level=1)

table = document.add_table(rows=2, cols=2, style = 'Medium Shading 1 Accent 3')

# Header row
cell = table.cell(0, 0)
cell.text = 'No.of rows'
cell_font = cell.paragraphs[0].runs[0].font
cell_font.size = Pt(11)
cell_font.bold = True

cell = table.cell(0, 1)
cell.text = 'No.of columns'
cell_font = cell.paragraphs[0].runs[0].font
cell_font.size = Pt(11)
cell_font.bold = True

# Values
cell = table.cell(1, 0)
cell.text = F'{df.shape[0] :,}'
cell_font = cell.paragraphs[0].runs[0].font
cell_font.size = Pt(11)
cell_font.bold = False

cell = table.cell(1, 1)
cell.text = F'{df.shape[1] :,}'
cell_font = cell.paragraphs[0].runs[0].font
cell_font.size = Pt(11)
cell_font.bold = False


# In[ ]:


# Page 4a
# document.add_page_break()
p = document.add_paragraph(' ')

# Heading 1
document.add_heading('Dataframe columns summary', level=1)

# Rehsape the column data type dataframe into form that can be printed in MS Word
data = round(data_qlt_df[['column_name','col_data_type', 'non_null_values', 'null_values', 'count']], 2)

# add a table to the end and create a reference variable
# extra row is so we can add the header row
table = document.add_table(data.shape[0]+1, data.shape[1], style='Medium Shading 1 Accent 3')

# add the header rows.
for j in range(data.shape[1]):

    #header row first two columns
    if j <= 1:
        cell = table.cell(0, j)
        cell.text = F'{data.columns[j]}'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
    else:
        cell = table.cell(0, j)
        cell.text = F'{data.columns[j]}'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
        cell.paragraphs[0].alignment= WD_ALIGN_PARAGRAPH.RIGHT
        
    
# add the rest of the data frame
for i in range(data.shape[0]):
    for j in range(data.shape[1]):
        if j <= 1:
            cell = table.cell(i+1, j)
            cell.text = F'{data.values[i,j]}'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = False            
        else:
            cell = table.cell(i+1, j)
            cell.text = F'{data.values[i,j] :,}'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = False  
            cell.paragraphs[0].alignment= WD_ALIGN_PARAGRAPH.RIGHT           


# # Word - Column memory usage profile

# In[ ]:


document.add_page_break()
document.add_heading('Memory Usage Profile', 0)


# In[ ]:


# Page 5
p = document.add_paragraph(' ')

# Heading 1
document.add_heading('Data file size on disk vs. dataset size in memory', level=1)

# Create table
table = document.add_table(rows=3, cols=2, style = 'Medium Shading 1 Accent 3')

# Add column headers
cell = table.cell(0,0)
cell.text = 'Description'
cell_font = cell.paragraphs[0].runs[0].font
cell_font.size = Pt(11)
cell_font.bold = True  

cell = table.cell(0,1)
cell.text = 'Size in MB'
cell_font = cell.paragraphs[0].runs[0].font
cell_font.size = Pt(11)
cell_font.bold = True  
cell.paragraphs[0].alignment= WD_ALIGN_PARAGRAPH.RIGHT 

# Add values : Value Line 1
cell = table.cell(1,0)
cell.text = 'Data file size on disk'
cell_font = cell.paragraphs[0].runs[0].font
cell_font.size = Pt(11)
cell_font.bold = False  

cell = table.cell(1,1)
cell.text = F'{round(file_size, 2)  :,.2f}'
cell_font = cell.paragraphs[0].runs[0].font
cell_font.size = Pt(11)
cell_font.bold = False  
cell.paragraphs[0].alignment= WD_ALIGN_PARAGRAPH.RIGHT 

# Add values : Value Line 2
cell = table.cell(2,0)
cell.text = 'Dataset size in memory'
cell_font = cell.paragraphs[0].runs[0].font
cell_font.size = Pt(11)
cell_font.bold = False  

cell = table.cell(2,1)
cell.text = F'{round(df_mem, 2)  :,.2f}'
cell_font = cell.paragraphs[0].runs[0].font
cell_font.size = Pt(11)
cell_font.bold = False  
cell.paragraphs[0].alignment= WD_ALIGN_PARAGRAPH.RIGHT 

# Memory increase
p = document.add_paragraph('')
p = document.add_paragraph('Dataset increase in memory :  ')
p.add_run(str(round(sz_increase*100, 2)) + '%').bold = True

# Add graph
document.add_picture('fig_df_tot_memory.png', height=Inches(3), width=Inches(3))


# In[ ]:


# Page 6
document.add_page_break()

# Heading 1
document.add_heading('Dataframe column types and size in memory', level=1)

# Rehsape the column data type dataframe into form that can be printed in MS Word
# Using .reset_index() will make the index a column
data = round(plt_dtype.reset_index(), 2)


# add a table to the end and create a reference variable
# extra row is so we can add the header row
table = document.add_table(data.shape[0]+1, data.shape[1], style = 'Medium Shading 1 Accent 3')

# add the header rows.
for j in range(data.shape[1]):
    #header row first first columns
    if j == 0:
        cell = table.cell(0, j)
        cell.text = F'{data.columns[j]}'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
    else:
        cell = table.cell(0, j)
        cell.text = F'{data.columns[j]}'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
        cell.paragraphs[0].alignment= WD_ALIGN_PARAGRAPH.RIGHT
        
    
# add the rest of the data frame
for i in range(data.shape[0]):
    for j in range(data.shape[1]):
        if j == 0:
            cell = table.cell(i+1, j)
            cell.text = F'{data.values[i,j]}'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = False            
        else:
            cell = table.cell(i+1, j)
            cell.text = F'{data.values[i,j] :,.2f}'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = False  
            cell.paragraphs[0].alignment= WD_ALIGN_PARAGRAPH.RIGHT           


p = document.add_paragraph(' ')
p = document.add_paragraph('"col_data_type" : Column data type')
p = document.add_paragraph('"dtype_count" : Number of oclumns in the dataset of the given data type')
p = document.add_paragraph('"dtype_total" : Total memory in MB for the given data type')
p = document.add_paragraph('"dtype_%_total_mem" : Percentage of the memory used by the given data type out of the total memory used by the dataset')

document.add_picture('fig_cols_memory.png', height=Inches(3), width=Inches(6))

p = document.add_paragraph('In a memory heavy datasets the above information can shed light into which data type you need to focus if you need to optimise the memory usage.')
p = document.add_paragraph('e.g. may be convert "object" datatype to "category" type if the cardinality is low or may be down cast "float64" to float16 or smaller.')
p = document.add_paragraph('These decision need further information on column cardinality and max/min values which are covered in the next few sections.')


# In[ ]:


# Page 7
document.add_page_break()

# Heading 1
document.add_heading('Memory used by "object" data type', level=1)


# Rehsape the column data type dataframe into form that can be printed in MS Word
# Using .reset_index() will make the index a column
data = round(cardn_df.sort_values("unique_values_count"), 2)


# add a table to the end and create a reference variable
# extra row is so we can add the header row
table = document.add_table(data.shape[0]+1, data.shape[1], style = 'Medium Shading 1 Accent 3')

# add the header rows.
for j in range(data.shape[1]):
    #header row first first columns
    if j == 0:
        cell = table.cell(0, j)
        cell.text = F'{data.columns[j]}'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
    else:
        cell = table.cell(0, j)
        cell.text = F'{data.columns[j]}'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
        cell.paragraphs[0].alignment= WD_ALIGN_PARAGRAPH.RIGHT
        
    
# add the rest of the data frame
for i in range(data.shape[0]):
    for j in range(data.shape[1]):
        if j == 0:
            cell = table.cell(i+1, j)
            cell.text = F'{data.values[i,j]}'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = False            
        else:
            cell = table.cell(i+1, j)
            cell.text = F'{data.values[i,j] :,.2f}'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = False  
            cell.paragraphs[0].alignment= WD_ALIGN_PARAGRAPH.RIGHT           


p = document.add_paragraph(' ')
p = document.add_paragraph('"column_name" : Name of the column in the dataframe')
p = document.add_paragraph('"col_memory" : Memory used by the given column')
p = document.add_paragraph('"%_of_dtype_mem" : Percentage of memory used by the given column out of memory used by the column data type')
p = document.add_paragraph('"%_of_total_memory" : Percentage of the memory used by the given column out of the total memory used by the dataset')
p = document.add_paragraph('"unique_values_count" : Count of the unique values for the given column')

document.add_picture('fig_object_cols_memory.png', height=Inches(3), width=Inches(6))

p = document.add_paragraph(' ')
p = document.add_paragraph("Analysing how many unique values an 'object' column has will be useful to detrminewhich columns are good candidates for *Categorical* data type. In combination with the total memory used by 'object'data type and each 'object' data type column, decisions can be made on converting them Category type.Object or string data type columns with low cardinality is suitable for Category type.")
p.add_run("The threshold of 'low cardinality' depends on the domain of the data and data usage patterns.").bold = True


# In[ ]:


# Page 8
document.add_page_break()

# Heading 1
document.add_heading('Memory used by "Non-Object" data type', level=1)

# Rehsape the column data type dataframe into form that can be printed in MS Word
# Using .reset_index() will make the index a column
data = round(num_cardn_df.sort_values("unique_values_count"), 2)

# add a table to the end and create a reference variable
# extra row is so we can add the header row
table = document.add_table(data.shape[0]+1, data.shape[1], style = 'Medium Shading 1 Accent 3')

# add the header rows.
for j in range(data.shape[1]):
    #header row first first columns
    if j == 0:
        cell = table.cell(0, j)
        cell.text = F'{data.columns[j]}'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
    else:
        cell = table.cell(0, j)
        cell.text = F'{data.columns[j]}'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
        cell.paragraphs[0].alignment= WD_ALIGN_PARAGRAPH.RIGHT
        
    
# add the rest of the data frame
for i in range(data.shape[0]):
    for j in range(data.shape[1]):
        if j == 0:
            cell = table.cell(i+1, j)
            cell.text = F'{data.values[i,j]}'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = False            
        else:
            cell = table.cell(i+1, j)
            cell.text = F'{data.values[i,j] :,.2f}'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = False  
            cell.paragraphs[0].alignment= WD_ALIGN_PARAGRAPH.RIGHT           


p = document.add_paragraph(' ')
p = document.add_paragraph('"column_name" : Name of the column in the dataframe')
p = document.add_paragraph('"col_memory" : Memory used by the given column')
p = document.add_paragraph('"%_of_dtype_mem" : Percentage of memory used by the given column out of memory used by the column data type')
p = document.add_paragraph('"%_of_total_memory" : Percentage of the memory used by the given column out of the total memory used by the dataset')

document.add_picture('fig_non_object_cols_memory.png', height=Inches(3), width=Inches(6))

p = document.add_paragraph(' ')
p = document.add_paragraph("By analysing the min and max values of the numeric columns decions can be made to downcast the data type to more memory efficient storage types.")


# In[ ]:


# Page 9
document.add_page_break()

# Heading 1
document.add_heading('Columns with non-null values less than ' + "{:,.2f}".format(threshold_perc*100) + '%', level=1)

p = document.add_paragraph('The columns should contain at least  ' + "{:,.0f}".format(col_vals_threshold) + '  (' + "{:,.2f}".format((col_vals_threshold/df.shape[0])*100) + '%) non-empty rows out of  '+ "{:,}".format(df.shape[0]) + ' rows to be considered useful.')
p = document.add_paragraph('The non-empty values threshold can be set using the threshold_perc variable in the code.')


# Rehsape the column data type dataframe into form that can be printed in MS Word
# Using .reset_index() will make the index a column
data = round(null_vals_df.sort_values("non_null_values"), 2)

# add a table to the end and create a reference variable
# extra row is so we can add the header row
table = document.add_table(data.shape[0]+1, data.shape[1], style = 'Medium Shading 1 Accent 3')

# add the header rows.
for j in range(data.shape[1]):
    #header row first first columns
    if j <= 1:
        cell = table.cell(0, j)
        cell.text = F'{data.columns[j]}'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
    else:
        cell = table.cell(0, j)
        cell.text = F'{data.columns[j]}'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
        cell.paragraphs[0].alignment= WD_ALIGN_PARAGRAPH.RIGHT
        
    
# add the rest of the data frame
for i in range(data.shape[0]):
    for j in range(data.shape[1]):
        if j <= 1:
            cell = table.cell(i+1, j)
            cell.text = F'{data.values[i,j]}'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = False            
        else:
            cell = table.cell(i+1, j)
            cell.text = F'{data.values[i,j] :,.2f}'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = False  
            cell.paragraphs[0].alignment= WD_ALIGN_PARAGRAPH.RIGHT           
        


p = document.add_paragraph(' ')
p = document.add_paragraph('"column_name" : Name of the column in the dataframe')
p = document.add_paragraph('"col_data_type" : Data type of the given column')
p = document.add_paragraph('"col_memory" : Memory used by the given column')
p = document.add_paragraph('"non_null_values" : Count of non-null values in the given column')
p = document.add_paragraph('"%_of_non_nulls" : Percentage of the non-null values out of total values for the given column')
p = document.add_paragraph('"null_values" : Count of null values in the given column')
p = document.add_paragraph('"%_of_nulls" : Percentage of the null values out of total values for the given column')

p = document.add_paragraph(' ')
p = document.add_paragraph("Generally columns with large percentage of empty values can be *dropped* from the dataset as they will not add any value to the analysis.")
p = document.add_paragraph('')
p.add_run('But this depends on the domian of the dataset and usage pattern of the columns/data.').bold = True


# # Word - Data Correlation plot

# In[ ]:


document.add_page_break()
document.add_heading('Data correlation plot', 0)

p = document.add_paragraph('')

document.add_picture('fig_cor_plot.png', height=Inches(6), width=Inches(6))


# # Word - Create the detail column profile rows

# In[ ]:


document.add_page_break()
document.add_heading('Column Data Profile Details', 0)


# In[ ]:


# ind = 1  # to be taken from iterrows loop later
for ind in range(data_qlt_df.shape[0]):
    document.add_page_break()
    
    # Create table for column profile details
    table = document.add_table(rows=6, cols=6, style = 'Medium Shading 1 Accent 3' )
    
    # Merge cells in header row for COlumn Name
    for y in range(len(table.rows[0].cells)-1):
        a = table.cell(0,y)
        b = table.cell(0,y+1)
        a.merge(b)

    # Merge cells in detail rows spanning 2 cells x 3 
    for row in range(1,6):
        a = table.cell(row,0)
        b = table.cell(row,1)
        a.merge(b)
        a = table.cell(row,2)
        b = table.cell(row,3)
        a.merge(b)
        a = table.cell(row,4)
        b = table.cell(row,5)
        a.merge(b)


    #*** ADD VALUES TO TABLE  ***#
    # Cell 0,0 (merged 6 cells): Header - Column Name
    cell = table.cell(0, 0)
    cell.text = data_qlt_df["column_name"][ind]
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(15)
    cell_font.bold = True

    # Cell 1,0: Blank
    cell = table.cell(1, 1)
    cell.text = "TBD Column :\n"
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run('no value')
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(12)
    cell_font2.bold = False

    # Cell 1,0: Column data type
    cell = table.cell(1, 3)
    cell.text = 'Data Type : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(str(data_qlt_df["col_data_type"][ind]))
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(12)
    cell_font2.bold = False

    # Cell 1,1: Count of toal values in the column
    cell = table.cell(1, 5)
    cell.text = 'Values Count : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["count"][ind] :,.0f}')
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False

    # Cell 2,0: Count of unique values in the column
    cell = table.cell(2, 1)
    cell.text = 'Unique Values Count : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    unique_per = (data_qlt_df["unique_values_count"][ind] / data_qlt_df["count"][ind]) * 100
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["unique_values_count"][ind] :,.0f}' + "   " + F'({unique_per :,.2f}%)' )
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False

    # Cell 2,1: Count of non-null values in the column
    cell = table.cell(2, 3)
    cell.text = 'Non-Null Values Count : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["non_null_values"][ind] :,.0f}' + "   " + F' ({data_qlt_df["%_of_non_nulls"][ind]  :,.2f}%)' )
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False       

    # Cell 2,2: Count of null values in the column
    cell = table.cell(2, 5)
    cell.text = 'Null Values Count : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["null_values"][ind]  :,.0f}' + "   " + F' ({data_qlt_df["%_of_nulls"][ind]  :,.2f}%)' )
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False

    # Cell 3,0: Min of values in the column
    cell = table.cell(3, 1)
    cell.text = 'Min : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["min"][ind]  :,.2f}' )
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False

    # Cell 3,1: Mean of values in the column
    cell = table.cell(3, 3)
    cell.text = 'Mean :  \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["mean"][ind] :,.2f}' )
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False

    # Cell 3,3: Max of values in the column
    cell = table.cell(3, 5)
    cell.text = 'Max : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["max"][ind]  :,.2f}' )
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False

    # Cell 4,1: 25th Percentile of values in the column
    cell = table.cell(4, 1)
    cell.text = '25th Percentile : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["25%"][ind]  :,.2f}' )
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False

    # Cell 4,2: 50th Percentile of values in the column
    cell = table.cell(4, 3)
    cell.text = '50th Percentile : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["50%"][ind]  :,.2f}' )
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False

    # Cell 4,3: 75th Percentile of values in the column
    cell = table.cell(4, 5)
    cell.text = '75th Percentile : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["75%"][ind]  :,.2f}' )
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False

    # Cell 5,1: Memory used by the column values
    cell = table.cell(5, 1)
    cell.text = 'Column Memory : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["col_memory"][ind] :,.2} MB' )
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False

    # Cell 5,2: Memory used by the column values vs. memory used by the data type
    cell = table.cell(5, 3)
    cell.text = 'As % of Dtype Memory  : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["%_of_dtype_mem"][ind] :.2f}%' )
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False                               

    # Cell 5,3: Memory used by the column values vs. memory used by the data type
    cell = table.cell(5, 5)
    cell.text = 'As % of DF Memory : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["%_of_total_memory"][ind] :.2f}%' )
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False

    p = document.add_paragraph(' ')
    p = document.add_paragraph(' ')

    fig_name = 'fig_' + data_qlt_df['column_name'][ind] + '.png'
    document.add_picture(fig_name, height=Inches(3.5), width=Inches(6))


# In[ ]:


# save the doc
document.save('data_profile_df_MS_WORD.docx')


# In[ ]:


print("Document generated!")


# In[ ]:




