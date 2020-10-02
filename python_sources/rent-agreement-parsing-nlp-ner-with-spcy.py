#!/usr/bin/env python
# coding: utf-8

# Meta Data Extraction,
# Named entity Recognition,
# Rent Agreement Parsing,
# Rent Agreement Meta Data Extraction,
# NLP,
# Spacy NER

# In[ ]:


#My library Templete
import numpy as np # linear algebra
import pandas as pd # data processing, CSV file I/O (e.g. pd.read_csv)
import os
import spacy
import tensorflow as tf
import numpy as np # linear algebra
import pandas as pd
import matplotlib.pyplot as plt
from sklearn.naive_bayes import GaussianNB
from matplotlib import pyplot
import gc
from tensorflow.keras import regularizers
from sklearn.preprocessing import minmax_scale
from sklearn.datasets import load_digits
from sklearn.model_selection import learning_curve as learning_cv
from sklearn.model_selection import ShuffleSplit
from sklearn.metrics import roc_curve, auc
from gensim.models import Word2Vec
from sklearn.decomposition import PCA
from sklearn.multiclass import OneVsRestClassifier
from scipy import interp
from sklearn.metrics import roc_auc_score
from sklearn.model_selection import validation_curve
import scikitplot as skplt   
from sklearn.metrics import classification_report   
import re
from tensorflow.keras.models import Sequential
from sklearn.metrics import roc_curve
from sklearn.metrics import auc
import seaborn as sns
from sklearn import metrics
import os
import json
from sklearn.preprocessing import minmax_scale
from sklearn.datasets import load_digits
from sklearn.model_selection import learning_curve as learning_cv
from sklearn.model_selection import ShuffleSplit
from sklearn.metrics import roc_curve, auc
from gensim.models import Word2Vec
from sklearn.decomposition import PCA
from sklearn.multiclass import OneVsRestClassifier
from scipy import interp
from sklearn.metrics import roc_auc_score
from sklearn.model_selection import validation_curve
import scikitplot as skplt   
from sklearn.metrics import classification_report   
import random
import logging
from sklearn.metrics import classification_report
from sklearn.metrics import precision_recall_fscore_support
from spacy.gold import GoldParse
from spacy.scorer import Scorer
from sklearn.metrics import accuracy_score


# In[ ]:


#install docx lib to read docx
get_ipython().system('pip install python-docx')


# In[ ]:


import pandas as pd
import docx
train_labels=pd.read_csv('/kaggle/input/rent-agreement/TrainingTestSet (2).csv')
val_labels=pd.read_csv('/kaggle/input/rent-agreement/ValidationSet.csv')
train_text_dir='/kaggle/input/rent-agreement/Training_data/Training_data/'
val_text_dir='/kaggle/input/rent-agreement/Validation_data/Validation_data/'


# In[ ]:


import re
def clean(text):
    text=text.replace('/','.')
    text=text.replace('\\','.')
    text=text.replace('. ','.')
    text=text.replace(', ','')
    text=text.replace(',','')
    text=text.replace('-','.')
    text=text.replace('/','.')
    text=text.replace('. ','.')
    #text=text.lower()
    return(text)


# In[ ]:


#prepair data for training
data=[]
data_index=list(train_labels.columns)
count=0
new_data=[]
for i in os.listdir('/kaggle/input/rent-agreement/Training_data/Training_data/'):
    doc = docx.Document('/kaggle/input/rent-agreement/Training_data/Training_data/'+i)  # Creating word reader object.
    paragraph=doc.paragraphs
    st=""
    for j in paragraph:
        st+=clean(str(j.text))+" "
    data.append(st)
    temp=train_labels.loc[train_labels['File Name'] == i[:-9]]
    
    entities=[]
    for col in data_index[1:]:
        query=str(tuple(temp[col])[0])

        if(col=='Aggrement Start Date' or col=='Aggrement End Date'):
            if(st.find(query)!=-1):
                entities.append([st.find(query),st.find(query)+len(query),col])
                continue

            
        
        if(type(tuple(temp[col])[0])==type(8.0) and tuple(temp[col])[0]>=0): 
                query=str(int(tuple(temp[col])[0]))      
        if(st.find(query)!=-1 or st.lower().find(query.lower())!=-1):
                entities.append([st.find(query),st.find(query)+len(query),col])

        else:
            #print(i,col,query,type(tuple(temp[col])[0]))
            count+=1
    new_data.append((st,{'entities':entities}))
print(count)
            


# In[ ]:


"""
data=[]
data_index=list(train_labels.columns)
count=0
for i in os.listdir('/kaggle/input/rent-agreement/Training_data/Training_data/'):
    doc = docx.Document('/kaggle/input/rent-agreement/Training_data/Training_data/'+i)  # Creating word reader object.
    paragraph=doc.paragraphs
    st=""
    for i in paragraph:
        st+=clean(i.text)+" "
    data.append(st)
    temp=train_labels.loc[train_labels['File Name'] == i]
    for col in data_index[1:]:
        query=tuple(temp[col])[0]
        if(st.lower().find(query.lower())!=-1):
            print(st.find(query),st.find(query)+len(query),st[st.find(query):st.find(query)+len(query)])
        else:
            print(i,query)
            count+=1
print(count)
"""            


# In[ ]:


TRAIN_DATA=new_data
#load model
nlp = spacy.blank('en')
if 'ner' not in nlp.pipe_names:
    ner = nlp.create_pipe('ner')
    nlp.add_pipe(ner, last=True)
#adding labels
for _, labels in TRAIN_DATA:
     for ent in labels.get('entities'):
        ner.add_label(ent[2])
#adding task to piplines
other_pipes = [pipe for pipe in nlp.pipe_names if pipe != 'ner']
count=0
loss_val=[]
with nlp.disable_pipes(*other_pipes):  # only train NER
    optimizer = nlp.begin_training()
    for i in range(100):
        print("iteration :" ,i," Start")
        random.shuffle(TRAIN_DATA)
        losses = {}
        #start training
        for text, annotations in TRAIN_DATA:
            try:                
                nlp.update(
                    [text],  # batch of texts
                    [annotations],  # batch of annotations
                    drop=0.2,  # dropout - make it harder to memorise data
                    sgd=optimizer,
                    # callable to update weights
                    losses=losses)
            except:
                pass
                count=count+1
        loss_val.append(list(losses.values())[0])
        print("Loss in iteration :",i,"loss : ",list(losses.values())[0])
#plot loss
plt.plot(loss_val)


# In[ ]:


#load validation data 
data=[]
data_index=list(train_labels.columns)
count=0
new_data=[]
for i in os.listdir('/kaggle/input/rent-agreement/Validation_Data/Validation_Data/'):
    doc = docx.Document('/kaggle/input/rent-agreement/Validation_Data/Validation_Data/'+i)  # Creating word reader object.
    paragraph=doc.paragraphs
    st=""
    for j in paragraph:
        st+=clean(str(j.text))+" "
    data.append(st)
    temp=val_labels.loc[val_labels['File Name'] == i[:-9]]
    #convert in proper format
    entities=[]
    for col in data_index[1:]:
        query=str(tuple(temp[col])[0])

        if(col=='Aggrement Start Date' or col=='Aggrement End Date'):
            if(st.find(query)!=-1):
                entities.append([st.find(query),st.find(query)+len(query),col])
                continue        
        if(type(tuple(temp[col])[0])==type(8.0) and tuple(temp[col])[0]>=0): 
                query=str(int(tuple(temp[col])[0]))      
        if(st.find(query)!=-1 or st.lower().find(query.lower())!=-1):
                entities.append([st.find(query),st.find(query)+len(query),col])

        else:
            print(i,col,query,type(tuple(temp[col])[0]))
            count+=1
    new_data.append((st,{'entities':entities}))


# In[ ]:


#test the model and evaluate it


# In[ ]:


#print predicted values
examples =new_data
c=0    
for text,annot in examples:
        doc_to_test=nlp(text)
        d={}
        for ent in doc_to_test.ents:
            d[ent.label_]=[]
        for ent in doc_to_test.ents:
            d[ent.label_].append(ent.text)
        print(d)
        c+=1

